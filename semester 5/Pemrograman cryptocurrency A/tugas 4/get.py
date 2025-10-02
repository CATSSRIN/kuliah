import requests
from bs4 import BeautifulSoup
from docx import Document

def scrape_and_save_to_docx(url, output_filename='hasil_scraping.docx'):
    """
    Fungsi untuk mengambil data dari URL dan menyimpannya ke file .docx.
    
    Args:
        url (str): URL dari halaman web yang ingin di-scrape.
        output_filename (str): Nama file .docx untuk menyimpan hasil.
    """
    try:
        # 1. Siapkan header untuk menyamar sebagai browser (INI BAGIAN PERBAIKAN)
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        # 2. Mengirim permintaan dengan header untuk menghindari blokir 403
        print(f"Mengambil data dari: {url}")
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # Akan menampilkan error jika request gagal

        # 3. Menggunakan BeautifulSoup untuk mem-parsing HTML
        soup = BeautifulSoup(response.text, 'html.parser')

        # 4. Mencari semua elemen yang berisi informasi artikel.
        # PENTING: Class HTML ('obj_article_summary', 'title', 'authors')
        # mungkin berbeda untuk website lain. Anda perlu memeriksanya
        # dengan fitur "Inspect" di browser.
        articles = soup.find_all('div', class_='obj_article_summary')
        
        if not articles:
            print("Peringatan: Tidak ada artikel yang ditemukan. Kemungkinan class HTML pada kode tidak cocok dengan website.")
            print("Silakan periksa kode sumber halaman web untuk menemukan class yang tepat.")
            return

        # 5. Membuat dokumen Word baru
        doc = Document()
        doc.add_heading('Hasil Pencarian Artikel', level=1)
        doc.add_paragraph(f'Sumber URL: {url}\n')

        print(f"Menemukan {len(articles)} artikel. Memproses...")
        # 6. Loop melalui setiap artikel dan ekstrak datanya
        for article in articles:
            title_element = article.find('h4', class_='title').find('a')
            title = title_element.get_text(strip=True) if title_element else "Judul tidak ditemukan"

            authors_element = article.find('div', class_='authors')
            authors = authors_element.get_text(strip=True) if authors_element else "Penulis tidak ditemukan"
            
            # 7. Menambahkan data ke dokumen Word
            doc.add_heading(title, level=3)
            doc.add_paragraph(f"Penulis: {authors}")
            
            # Menambahkan link sumber jika ditemukan
            if title_element and 'href' in title_element.attrs:
                doc.add_paragraph(f"Link: {title_element['href']}")

            doc.add_paragraph() # Menambah spasi antar artikel

        # 8. Menyimpan dokumen
        doc.save(output_filename)
        print(f"Sukses! Data telah disimpan ke file '{output_filename}'")

    except requests.exceptions.HTTPError as e:
        print(f"Gagal mengakses URL. Status Code: {e.response.status_code}")
        print("Website mungkin memblokir akses. Coba ganti User-Agent atau gunakan VPN.")
    except requests.exceptions.RequestException as e:
        print(f"Gagal mengakses URL: {e}")
    except Exception as e:
        print(f"Terjadi kesalahan: {e}")


# --- CARA MENGGUNAKAN ---
if __name__ == '__main__':
    # Ganti URL ini dengan link yang Anda inginkan
    target_url = "https://jurnal.untan.ac.id/index.php/jepin/search/search?simpleQuery=prediksi&searchField=query"
    
    # Panggil fungsi untuk memulai proses
    scrape_and_save_to_docx(target_url, output_filename='hasil_jurnal_untan.docx')