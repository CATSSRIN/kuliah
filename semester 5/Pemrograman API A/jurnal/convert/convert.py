from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_journal_doc():
    doc = Document()
    
    # --- STYLE SETUP (Sederhana) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10) # Asumsi ukuran font standar jurnal

    # --- JUDUL  ---
    title = doc.add_paragraph("Pengembangan RESTful API Cuaca Terintegrasi Dengan BMKG Untuk Sistem Informasi Real-Time Berbasis Cloud")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # --- PENULIS  ---
    authors = doc.add_paragraph("Peneliti Utama, Peneliti Kedua, Peneliti Ketiga")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.runs[0].bold = True
    
    # --- AFILIASI  ---
    affil = doc.add_paragraph("Program Studi Teknik Informatika, Fakultas Teknologi Informasi, Universitas Negeri Indonesia")
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil.runs[0].italic = True
    
    # --- TANGGAL  ---
    dates = doc.add_paragraph("Diterima: 11 Januari 2025 | Revisi: 15 Maret 2025 | Diterbitkan: 20 Juni 2025")
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer

    # --- ABSTRAK INDONESIA [cite: 6-13] ---
    doc.add_heading('ABSTRAK', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs = doc.add_paragraph("Abstrak artikel ini terdiri dari 150 kata yang mencakup: (1) Tujuan Utama: Penelitian ini bertujuan mengembangkan sistem RESTful API yang terintegrasi dengan data cuaca real-time dari BMKG. (2) Kebaruan: Implementasi caching database lokal dengan integrasi langsung. (3) Metode Penelitian: Rapid Application Development (RAD) dengan sampling 10 kota. (4) Temuan: Response time rata-rata 145ms, uptime 99.8%. (5) Kesimpulan: API terbukti efektif untuk integrasi data cuaca BMKG.")
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- KATA KUNCI  ---
    doc.add_paragraph("Kata Kunci: RESTful API, BMKG, Cuaca Real-time, PHP, Caching")
    
    # --- ABSTRACT INGGRIS [cite: 15-24] ---
    doc.add_heading('ABSTRACT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs_en = doc.add_paragraph("This article abstract consists of 150 words covering: (1) Main Purpose: To develop a RESTful API integrated with BMKG. (2) Novelty: Caching implementation. (3) Methods: RAD. (4) Findings: 145ms response time. (5) Conclusion: Effective integration.")
    p_abs_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("Keywords: RESTful API, BMKG, Real-time Weather, PHP, Caching")
    
    # --- HOW TO CITE [cite: 25-27] ---
    doc.add_paragraph("How to Cite:")
    cite_text = doc.add_paragraph("Peneliti Utama, P., Peneliti Kedua, P., & Peneliti Ketiga, P. (2025). Pengembangan RESTful API Cuaca Terintegrasi dengan BMKG untuk Sistem Informasi Real-Time Berbasis Cloud. Scan : Jurnal Teknologi Informasi dan Komunikasi, 15(2), 145-162.")
    cite_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- PENDAHULUAN  ---
    doc.add_heading('PENDAHULUAN', level=1)
    p_intro = doc.add_paragraph("Perkembangan teknologi informasi di era digital telah mengubah cara masyarakat mengakses informasi meteorologi. [Salin isi lengkap Latar Belakang dan Tinjauan Pustaka dari MD Anda di sini, gabungkan menjadi satu alur cerita].")
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- METODE PENELITIAN  ---
    doc.add_heading('METODE PENELITIAN', level=1)
    p_met = doc.add_paragraph("Penelitian ini menggunakan pendekatan Rapid Application Development (RAD). [Salin isi Metode dari MD di sini].")
    p_met.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- HASIL DAN PEMBAHASAN  ---
    doc.add_heading('HASIL DAN PEMBAHASAN', level=1)
    p_res = doc.add_paragraph("Sistem RESTful API Cuaca BMKG yang dikembangkan mengadopsi arsitektur berlapis. [Salin isi Hasil dari MD di sini].")
    p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Contoh Tabel
    doc.add_paragraph("Tabel 1. Hasil Pengujian Throughput")
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    # Isi header tabel manual sesuai kebutuhan
    
    doc.add_paragraph("Sumber: Hasil Pengolahan Data Peneliti (2025)")

    # --- SIMPULAN  ---
    doc.add_heading('SIMPULAN', level=1)
    p_con = doc.add_paragraph("Penelitian ini berhasil mengembangkan RESTful API Cuaca BMKG yang terintegrasi, scalable, dan reliable. [Salin isi Simpulan dari MD].")
    p_con.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- DAFTAR PUSTAKA  ---
    doc.add_heading('DAFTAR PUSTAKA', level=1)
    p_ref = doc.add_paragraph("[Salin daftar pustaka dari MD di sini]")
    
    doc.save('Jurnal_Formatted.docx')
    print("File Jurnal_Formatted.docx berhasil dibuat!")

if __name__ == "__main__":
    create_journal_doc()